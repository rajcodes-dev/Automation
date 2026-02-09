import docx
import os

guest_file = 'guests.txt'
template_file = 'invitation_template.docx'
output_file = 'all_invitations.docx'

if not os.path.exists(template_file):
    print(f"Error: {template_file} not found.")
else:
    doc = docx.Document(template_file)
    with open(guest_file, 'r') as f:
        guests = f.readlines()

    for guest in guests:
        guest = guest.strip()

        if not guest:
            continue

        doc.add_paragraph("It would be a pleasure to have the company of", style='Normal')

        doc.add_paragraph(guest, style='Heading 1')
        doc.add_paragraph('at 11010 Memory Lane on the Evening of', style='Normal')
        doc.add_paragraph('April 1st', style='Normal')
        doc.add_paragraph('at 7 o\'clock', style='Normal')

        doc.add_page_break()

    doc.save(output_file)
    print(f"Success! Invitation saved to {output_file}")
